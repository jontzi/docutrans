import os
from docx import Document
from deep_translator import ChatGptTranslator
import time
from tqdm import tqdm
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Get API key from environment
api_key = os.getenv('OPENAI_API_KEY')

# Debug mode flag - set to False for full translation
DEBUG_MODE = False
DEBUG_PAGES = 5

# Ensure the API key is available
if not api_key:
    raise ValueError("Please set the OPENAI_API_KEY environment variable.")

# Initialize the ChatGPT Translator from deep-translator
translator = ChatGptTranslator(api_key=api_key, source='en', target='fi')

# Load the DOCX file
input_file = r"C:\Users\L11499\OneDrive - OPPO\translate\input\ColorOS_14.0_User_Manual_V1.0_20231227.docx"  # Replace with your input file name
output_file = r"C:\Users\L11499\OneDrive - OPPO\translate\output\ColorOS_14.0_User_Manual_V1.0_20231227_translated.docx"

# Open and read the document
document = Document(input_file)
translated_document = Document(input_file)  # Open source document to preserve formatting

# Find page breaks to determine page boundaries
page_breaks = []
current_page = 1
for i, para in enumerate(document.paragraphs):
    for run in para.runs:
        if len(run._r.findall('.//w:br[@w:type="page"]', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})) > 0:
            page_breaks.append(i)
            current_page += 1
            break

# Determine the paragraph index for the debug page limit
debug_end_para = len(document.paragraphs)
if DEBUG_MODE and page_breaks:
    for i, break_idx in enumerate(page_breaks):
        if i + 1 >= DEBUG_PAGES:
            debug_end_para = break_idx
            break

# Add counter variables before translation starts
total_paragraphs = len([p for p in document.paragraphs[:debug_end_para] if p.text.strip()])
total_cells = sum(len(row.cells) for table in document.tables for row in table.rows)
start_time = time.time()

if DEBUG_MODE:
    print(f"DEBUG MODE: Translating only first {DEBUG_PAGES} pages")

# Add temp directory path
temp_dir = r"translate\temp"
os.makedirs(temp_dir, exist_ok=True)

# Translating paragraphs
print(f"Translating {total_paragraphs} paragraphs...")
translated_count = 0
for i, para in tqdm(enumerate(document.paragraphs), total=len(document.paragraphs[:debug_end_para]), 
                   desc="Paragraphs", unit="para", ncols=100):
    if DEBUG_MODE and i >= debug_end_para:
        break
    if para.text.strip():  # Skip empty paragraphs
        translated_text = translator.translate(para.text)
        # Replace text while keeping the formatting
        translated_document.paragraphs[i].text = translated_text
        translated_count += 1
        if translated_count % 10 == 0:
            elapsed_time = time.time() - start_time
            speed = translated_count / elapsed_time
            # Save temp file
            temp_file = os.path.join(temp_dir, f"temp_translation_{translated_count}.docx")
            translated_document.save(temp_file)

# Remove paragraphs after debug limit
if DEBUG_MODE:
    while len(translated_document.paragraphs) > debug_end_para:
        p = translated_document.paragraphs[-1]._element
        p.getparent().remove(p)

# Handling tables (if any)
print(f"\nTranslating table cells...")
table_cells_translated = 0
for table_index, table in enumerate(document.tables):
    if DEBUG_MODE and table_index >= DEBUG_PAGES:
        break
    for row_index, row in enumerate(tqdm(table.rows, total=len(table.rows), 
                                       desc=f"Table {table_index+1}", unit="row", ncols=100)):
        for col_index, cell in enumerate(row.cells):
            # Get all paragraphs in the cell and translate each one
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    translated_text = translator.translate(paragraph.text)
                    # Clear existing paragraphs and add translated text
                    if len(cell.paragraphs) == 1:
                        cell.text = translated_text
                    else:
                        paragraph.text = translated_text
                    table_cells_translated += 1
                    if table_cells_translated % 10 == 0:
                        elapsed_time = time.time() - start_time
                        speed = (translated_count + table_cells_translated) / elapsed_time
                        print(f"\nSpeed: {speed:.2f} translations/second")
                        # Save temp file
                        temp_file = os.path.join(temp_dir, f"temp_translation_p{translated_count}_t{table_cells_translated}.docx")
                        translated_document.save(temp_file)
                        print(f"Saved temporary file: {temp_file}")

# Remove tables after debug limit
if DEBUG_MODE:
    while len(translated_document.tables) > DEBUG_PAGES:
        tbl = translated_document.tables[-1]._tbl
        tbl.getparent().remove(tbl)

# Calculate and display final statistics
total_time = time.time() - start_time
total_translations = translated_count + table_cells_translated
final_speed = total_translations / total_time

print(f"\nTranslation completed in {total_time:.2f} seconds")
print(f"Total items translated: {total_translations}")
print(f"Average speed: {final_speed:.2f} translations/second")

# Save the translated document
translated_document.save(output_file)

print(f"The translated file is saved as '{output_file}'.")
