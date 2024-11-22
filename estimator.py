from openai import OpenAI
from docx import Document
import tiktoken

def count_tokens(text, model="gpt-3.5-turbo"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

# Load the DOCX file
input_file = r"C:\Users\L11499\OneDrive - OPPO\translate\input\ColorOS_14.0_User_Manual_V1.0_20231227.docx"  # Replace with your input file name
# Load the document
document = Document(input_file)
total_tokens = 0
total_text_elements = 0
non_empty_text_elements = 0

# Count paragraphs
for para in document.paragraphs:
    total_text_elements += 1
    if para.text.strip():
        tokens = count_tokens(para.text)
        total_tokens += tokens
        total_tokens += 50
        non_empty_text_elements += 1

# Count table cells
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            total_text_elements += 1
            if cell.text.strip():
                tokens = count_tokens(cell.text)
                total_tokens += tokens
                total_tokens += 50
                non_empty_text_elements += 1

print(f"Total number of text elements (including empty): {total_text_elements}")
print(f"Number of non-empty text elements: {non_empty_text_elements}")
print(f"Estimated total tokens needed: {total_tokens}")
print(f"Estimated cost (USD): ${(total_tokens/1000) * 0.002:.2f}")  # $0.002 per 1K tokens for gpt-3.5-turbo 